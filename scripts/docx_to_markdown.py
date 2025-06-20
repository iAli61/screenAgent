#!/usr/bin/env python3
"""
DOCX to Markdown Converter

This script converts Microsoft Word (.docx) files to Markdown (.md) format.
"""

import os
import re
import sys
import argparse
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def extract_text_with_formatting(paragraph):
    """Extract text from a paragraph while preserving basic formatting."""
    text = ""
    for run in paragraph.runs:
        if run.text.strip():
            # Apply formatting
            run_text = run.text
            if run.bold:
                run_text = f"**{run_text}**"
            if run.italic:
                run_text = f"*{run_text}*"
            if run.underline:
                run_text = f"__{run_text}__"
            text += run_text
    return text


def convert_docx_to_markdown(docx_path, output_path=None):
    """
    Convert a .docx file to markdown format.
    
    Args:
        docx_path (str): Path to the .docx file
        output_path (str, optional): Path where the markdown file will be saved.
                                    If not provided, it will use the same name as the docx file.
    
    Returns:
        str: Path to the created markdown file
    """
    try:
        # Load the document
        doc = Document(docx_path)
        
        # Determine output path if not provided
        if output_path is None:
            base_name = os.path.splitext(docx_path)[0]
            output_path = f"{base_name}.md"
        
        markdown_lines = []
        
        # Process document paragraphs
        for para in doc.paragraphs:
            if not para.text.strip():
                # Add an empty line for empty paragraphs
                markdown_lines.append("")
                continue
                
            # Handle headings based on style
            style_name = para.style.name.lower()
            text = extract_text_with_formatting(para)
            
            if 'heading' in style_name:
                # Extract heading level (Heading 1, Heading 2, etc.)
                heading_level = ''
                for char in style_name:
                    if char.isdigit():
                        heading_level = char
                        break
                
                if heading_level:
                    level = int(heading_level)
                    if 1 <= level <= 6:  # Valid markdown heading levels
                        markdown_lines.append(f"{'#' * level} {text}")
                        continue
            
            # Handle alignment
            if para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                markdown_lines.append(f"<center>{text}</center>")
            elif para.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                markdown_lines.append(f"<div align='right'>{text}</div>")
            else:
                markdown_lines.append(text)
        
        # Handle tables
        for table in doc.tables:
            markdown_lines.append("\n")  # Add space before table
            
            # Process each row in the table
            for i, row in enumerate(table.rows):
                # Create the table row in markdown
                row_content = []
                for cell in row.cells:
                    # Get cell text, combine paragraphs
                    cell_text = " ".join([p.text for p in cell.paragraphs if p.text.strip()])
                    # Replace | character as it has special meaning in markdown tables
                    cell_text = cell_text.replace("|", "\\|")
                    row_content.append(cell_text)
                
                markdown_lines.append(f"| {' | '.join(row_content)} |")
                
                # After the first row, add the separator row for markdown tables
                if i == 0:
                    separator = "| " + " | ".join(["---"] * len(row.cells)) + " |"
                    markdown_lines.append(separator)
        
        # Write the markdown content to file
        with open(output_path, 'w', encoding='utf-8') as md_file:
            md_file.write('\n'.join(markdown_lines))
        
        return output_path
    
    except Exception as e:
        print(f"Error converting {docx_path} to markdown: {str(e)}")
        return None


def main():
    parser = argparse.ArgumentParser(description='Convert DOCX files to Markdown format.')
    parser.add_argument('input_file', help='Input DOCX file path')
    parser.add_argument('-o', '--output', help='Output Markdown file path (optional)')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input_file):
        print(f"Error: Input file '{args.input_file}' not found.")
        sys.exit(1)
    
    output_path = convert_docx_to_markdown(args.input_file, args.output)
    
    if output_path:
        print(f"Successfully converted '{args.input_file}' to '{output_path}'")
    else:
        print("Conversion failed.")
        sys.exit(1)


if __name__ == "__main__":
    main()
